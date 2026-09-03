import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_JUSTIFY


class Exporter:
    """Экспорт транскриптов в различные форматы"""

    def __init__(self):
        self.exports_dir = "exports"
        os.makedirs(self.exports_dir, exist_ok=True)

        # Регистрируем шрифт DejaVu Sans для поддержки кириллицы в PDF
        try:
            font_path = Path(__file__).parent.parent.parent / "fonts" / "DejaVuSans.ttf"
            print(f"Попытка загрузить шрифт из: {font_path}")
            print(f"Файл существует: {font_path.exists()}")

            if font_path.exists():
                pdfmetrics.registerFont(TTFont('DejaVuSans', str(font_path)))
                self.pdf_font = 'DejaVuSans'
                print(f"✅ Шрифт DejaVu Sans успешно зарегистрирован")
            else:
                print(f"❌ Файл шрифта не найден: {font_path}")
                self.pdf_font = 'Helvetica'  # Fallback
        except Exception as e:
            print(f"❌ Ошибка при загрузке шрифта DejaVu Sans: {e}")
            import traceback
            traceback.print_exc()
            self.pdf_font = 'Helvetica'

    def export_to_txt(self, text: str, filename: str = None) -> str:
        """
        Экспорт в TXT файл

        Args:
            text: Текст транскрипта
            filename: Имя файла (опционально)

        Returns:
            Путь к созданному файлу
        """
        if not filename:
            filename = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        filepath = os.path.join(self.exports_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(text)

        return filepath

    def export_to_docx(self, text: str, filename: str = None) -> str:
        """
        Экспорт в DOCX файл

        Args:
            text: Текст транскрипта
            filename: Имя файла (опционально)

        Returns:
            Путь к созданному файлу
        """
        if not filename:
            filename = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        filepath = os.path.join(self.exports_dir, filename)

        # Создаём документ
        doc = Document()

        # Добавляем заголовок
        heading = doc.add_heading('Транскрипт', level=1)

        # Добавляем дату
        date_para = doc.add_paragraph()
        date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True

        # Добавляем пустую строку
        doc.add_paragraph()

        # Добавляем текст транскрипта
        para = doc.add_paragraph(text)
        para_format = para.paragraph_format
        para_format.line_spacing = 1.5

        # Сохраняем
        doc.save(filepath)

        return filepath

    def export_to_pdf(self, text: str, filename: str = None) -> str:
        """
        Экспорт в PDF файл

        Args:
            text: Текст транскрипта
            filename: Имя файла (опционально)

        Returns:
            Путь к созданному файлу
        """
        if not filename:
            filename = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        filepath = os.path.join(self.exports_dir, filename)

        # Создаём PDF документ
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Контейнер для элементов
        story = []

        # Регистрируем Unicode шрифт прямо здесь
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.pdfbase import pdfmetrics
            pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
            font_name = 'HeiseiMin-W3'
        except:
            try:
                # Пробуем DejaVu Sans из файла
                font_path = Path(__file__).parent.parent.parent / "fonts" / "DejaVuSans.ttf"
                if font_path.exists():
                    pdfmetrics.registerFont(TTFont('DejaVuSans', str(font_path)))
                    font_name = 'DejaVuSans'
                else:
                    font_name = 'Helvetica'
            except:
                font_name = 'Helvetica'

        # Стили
        styles = getSampleStyleSheet()

        # Создаём кастомный стиль для русского текста
        style_normal = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
        )

        style_heading = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=16,
            leading=22,
        )

        style_date = ParagraphStyle(
            'CustomDate',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10,
            leading=14,
        )

        # Добавляем заголовок
        story.append(Paragraph("Транскрипт", style_heading))
        story.append(Spacer(1, 12))

        # Добавляем дату
        date_text = f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        story.append(Paragraph(date_text, style_date))
        story.append(Spacer(1, 24))

        # Добавляем текст транскрипта
        # Разбиваем на параграфы для лучшего форматирования
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Экранируем специальные символы для XML
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(para_text.replace('\n', '<br/>'), style_normal))
                story.append(Spacer(1, 12))

        # Строим документ
        doc.build(story)

        return filepath

    def cleanup_old_files(self, max_age_hours: int = 24):
        """
        Удаляет старые экспортированные файлы

        Args:
            max_age_hours: Максимальный возраст файлов в часах
        """
        import time

        now = time.time()
        max_age_seconds = max_age_hours * 3600

        for filename in os.listdir(self.exports_dir):
            filepath = os.path.join(self.exports_dir, filename)

            if os.path.isfile(filepath):
                file_age = now - os.path.getmtime(filepath)

                if file_age > max_age_seconds:
                    try:
                        os.remove(filepath)
                    except Exception:
                        pass
